"""
Document Parser
Extracts and normalizes text from PDFs, DOCX, and TXT files.
"""

import re
from pathlib import Path

import structlog
import pdfplumber
from docx import Document as DocxDocument

logger = structlog.get_logger(__name__)


class DocumentParser:
    """Extracts text content from various document formats."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt"}

    def parse(self, file_path: str | Path) -> str:
        """
        Parse a document and return its text content.
        
        Args:
            file_path: Path to the document file.
            
        Returns:
            Extracted and cleaned text content.
            
        Raises:
            ValueError: If file type is unsupported.
            FileNotFoundError: If file doesn't exist.
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}")

        logger.info("Parsing document", file=path.name, type=ext)

        extractors = {
            ".pdf": self._extract_pdf,
            ".docx": self._extract_docx,
            ".doc": self._extract_docx,
            ".txt": self._extract_txt,
        }

        raw_text = extractors[ext](path)
        cleaned = self._clean_text(raw_text)

        logger.info(
            "Parsed document",
            file=path.name,
            raw_chars=len(raw_text),
            cleaned_chars=len(cleaned),
        )
        return cleaned

    def _extract_pdf(self, path: Path) -> str:
        """Extract text from PDF using pdfplumber for better accuracy."""
        pages = []
        try:
            with pdfplumber.open(path) as pdf:
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text:
                        pages.append(text)
                    
                    # Also extract tables as text
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            if row:
                                row_text = " | ".join(
                                    str(cell) if cell else "" for cell in row
                                )
                                pages.append(row_text)

        except Exception as e:
            logger.error("PDF extraction failed, trying PyPDF2 fallback", error=str(e))
            return self._extract_pdf_fallback(path)

        return "\n\n".join(pages)

    def _extract_pdf_fallback(self, path: Path) -> str:
        """Fallback PDF extraction using PyPDF2."""
        from PyPDF2 import PdfReader

        reader = PdfReader(str(path))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n\n".join(pages)

    def _extract_docx(self, path: Path) -> str:
        """Extract text from DOCX files including paragraphs and tables."""
        doc = DocxDocument(str(path))
        content_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                content_parts.append(text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    content_parts.append(row_text)

        return "\n\n".join(content_parts)

    def _extract_txt(self, path: Path) -> str:
        """Extract text from plain text files."""
        encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]
        
        for encoding in encodings:
            try:
                return path.read_text(encoding=encoding)
            except (UnicodeDecodeError, UnicodeError):
                continue

        # Last resort: read with error replacement
        return path.read_text(encoding="utf-8", errors="replace")

    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text.
        
        - Remove excessive whitespace
        - Normalize line breaks
        - Remove non-printable characters
        - Fix common OCR artifacts
        """
        if not text:
            return ""

        # Remove non-printable characters (keep newlines, tabs)
        text = re.sub(r"[^\S\n\t]+", " ", text)
        
        # Remove control characters except newline and tab
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]", "", text)

        # Normalize multiple newlines to max 2
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Remove leading/trailing whitespace from each line
        lines = [line.strip() for line in text.split("\n")]
        text = "\n".join(lines)

        # Remove leading/trailing whitespace from full text
        text = text.strip()

        # Fix common OCR artifacts
        text = re.sub(r"(?<=[a-z])- ?\n(?=[a-z])", "", text)  # Hyphenated line breaks
        text = re.sub(r"\s{2,}", " ", text)  # Multiple spaces (but preserve newlines)

        # Re-normalize after fixes
        text = re.sub(r"\n{3,}", "\n\n", text)

        return text
